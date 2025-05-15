import os
import re
import json
import logging
import subprocess
from datetime import datetime
from dateutil import parser
from pdfminer.high_level import extract_text
from docx import Document
import google.generativeai as genai
from django.conf import settings

logger = logging.getLogger(__name__)
genai.configure(api_key=settings.GOOGLE_API_KEY)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc"}


def is_allowed_file(filename):
    return os.path.splitext(filename)[1].lower() in ALLOWED_EXTENSIONS


def extract_resume_text(file_path):
    try:
        ext = file_path.lower()
        if ext.endswith(".pdf"):
            return extract_text(file_path)
        elif ext.endswith(".docx"):
            doc = Document(file_path)
            return "\n".join(para.text for para in doc.paragraphs)
        elif ext.endswith(".doc"):
            temp_docx_path = file_path.replace(".doc", ".docx")
            subprocess.run(
                ["unoconv", "-f", "docx", "-o", temp_docx_path, file_path], check=True
            )
            doc = Document(temp_docx_path)
            return "\n".join(para.text for para in doc.paragraphs)
    except Exception as e:
        logger.error(f"Failed to extract text from {file_path}: {str(e)}")
    return ""


def parse_resume_with_gemini(resume_texts):
    prompt = (
        "You are an expert resume parser. Extract the following details for EACH resume:\n"
        "1. Name (full name exactly as shown)\n"
        "2. Email (complete address without spaces)\n"
        "3. Phone Number (with country code if available)\n"
        "4. Extract experience as a list of job roles with start_date and end_date (e.g., 'February 2021', 'Present'). No need to calculate years/months.\n"
        "5. Current Company Name (official legal name)\n"
        "6. Current Designation (exact job title)\n\n"
        "Return STRICT JSON array. Each object MUST follow this example:\n"
        "[\n"
        "  {\n"
        '    "name": "John Doe",\n'
        '    "email": "john@email.com",\n'
        '    "phoneNumber": "+11234567890",\n'
        '    "experiences": [\n'
        "      {\n"
        '        "job_title": "Software Engineer",\n'
        '        "company": "XYZ Ltd",\n'
        '        "start_date": "February 2021",\n'
        '        "end_date": "Present"\n'
        "      },\n"
        "      {\n"
        '        "job_title": "Intern",\n'
        '        "company": "ABC Corp",\n'
        '        "start_date": "July 2018",\n'
        '        "end_date": "January 2021"\n'
        "      }\n"
        "    ],\n"
        '    "currentCompanyName": "Tech Corp",\n'
        '    "currentDesignation": "Software Engineer"\n'
        "  }\n"
        "]\n\n"
        "Important Rules:\n"
        "- Phone numbers must start with '+' followed by country code\n"
        "- Remove all spaces from emails\n"
        "- Use full month names (January, February etc.)\n"
        "- If information is missing, use empty string\n"
        "- Current company is the most recent/last mentioned job\n"
        "- Return object should be proper JSON array of objects\n\n"
        "Resumes:\n"
        + "\n---\n".join(
            f"RESUME {i+1}:\n{text}" for i, text in enumerate(resume_texts)
        )
    )

    try:
        model = genai.GenerativeModel("gemini-2.0-flash-thinking-exp-01-21")
        response = model.generate_content(prompt)
        raw_response = response.text.strip()

        json_start = raw_response.find("[")
        json_end = raw_response.rfind("]") + 1
        json_str = raw_response[json_start:json_end] if json_start != -1 else "{}"
        if json_str.startswith("```json"):
            json_str = json_str.strip("```")[4:].strip()

        return json.loads(json_str)
    except Exception as e:
        logger.error(f"Gemini parsing failed: {str(e)}")
        logger.debug(f"Raw response: {raw_response}")
        return []


def calculate_experience(experiences):
    total_months = 0
    for exp in experiences:
        try:
            start = parser.parse(exp["start_date"], fuzzy=True)
            end_str = exp["end_date"]
            end = (
                datetime.now()
                if "present" in end_str.lower()
                else parser.parse(end_str, fuzzy=True)
            )
            months = (end.year - start.year) * 12 + (end.month - start.month)
            if months > 0:
                total_months += months
        except Exception:
            continue

    years = total_months // 12
    months = total_months % 12
    return {"year": years, "month": months}


def normalize_phone(number):
    if not number:
        return ""
    number = re.sub(r"[^\d+]", "", number)
    return number if number.startswith("+") else f"+{number}"


def process_resumes(file_paths):
    resume_texts, file_names = [], []
    for path in file_paths:
        if not is_allowed_file(path):
            continue
        text = extract_resume_text(path)
        if text:
            resume_texts.append(text)
            file_names.append(os.path.basename(path))
        else:
            logger.warning(f"No text extracted from {path}")

    parsed = parse_resume_with_gemini(resume_texts)
    results = []
    for i, data in enumerate(parsed):
        exp = calculate_experience(data.get("experiences", []))
        results.append(
            {
                "file_name": file_names[i],
                "name": data.get("name", "").strip(),
                "email": data.get("email", "").replace(" ", ""),
                "phone_number": normalize_phone(data.get("phoneNumber", "")),
                "years_of_experience": exp,
                "current_company": data.get("currentCompanyName", ""),
                "current_designation": data.get("currentDesignation", ""),
            }
        )
    return results
